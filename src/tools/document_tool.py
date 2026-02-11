from fpdf import FPDF
import uuid
import os
import openpyxl
try:
    from docx import Document
except ImportError:
    Document = None  # Handle case where python-docx isn't installed

class DocumentTool:
    def __init__(self):
        if not os.path.exists("data"):
            os.makedirs("data")

    def _sanitize_text(self, text):
        """
        Nuclear option for FPDF compatibility.
        Replaces smart quotes and non-latin characters that crash FPDF.
        """
        replacements = {
            '\u2018': "'",  # Left single quote
            '\u2019': "'",  # Right single quote
            '\u201c': '"',  # Left double quote
            '\u201d': '"',  # Right double quote
            '\u2013': '-',  # En dash
            '\u2014': '--', # Em dash
            '\u2026': '...', # Ellipsis
            '\u2022': '*',   # Bullet
        }
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        
        # Final fallback: force into latin-1, replacing unknown chars with '?'
        return text.encode('latin-1', 'replace').decode('latin-1')

    def create_pdf(self, content, filename=None):
        """Generates a PDF with robust encoding handling."""
        if not filename:
            filename = f"doc_{uuid.uuid4().hex[:8]}.pdf"
            
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=11)
        
        # 1. SANITIZE EVERYTHING FIRST
        content = self._sanitize_text(content)
        
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                pdf.ln(3)
                continue

            try:
                # CASE 1: Main Title (#)
                if line.startswith('# '):
                    pdf.set_font("Arial", 'B', 16)
                    pdf.multi_cell(0, 10, line.replace('# ', '').strip())
                    pdf.ln(2)

                # CASE 2: Section Heading (## or ###)
                elif line.startswith('##') or line.startswith('###'):
                    pdf.set_font("Arial", 'B', 12)
                    clean_text = line.replace('### ', '').replace('## ', '').strip()
                    pdf.multi_cell(0, 8, clean_text)
                    pdf.ln(1)
                
                # CASE 3: Bullet Points (* or -)
                elif line.startswith('* ') or line.startswith('- '):
                    pdf.set_font("Arial", '', 11)
                    clean_text = line[2:].strip()
                    current_x = pdf.get_x()
                    pdf.set_x(current_x + 5) 
                    pdf.multi_cell(0, 6, f"{chr(149)} {clean_text}")
                    pdf.set_x(current_x)

                # CASE 4: Standard Text
                else:
                    pdf.set_font("Arial", '', 11)
                    pdf.multi_cell(0, 6, line)
                    
            except Exception as e:
                print(f"[PDF ERROR] Line skipped: {e}")
                # Emergency fallback
                pdf.set_font("Arial", '', 11)
                pdf.multi_cell(0, 6, line)

        path = f"data/{filename}"
        try:
            pdf.output(path)
            return path
        except Exception as e:
            print(f"[CRITICAL PDF FAILURE] {e}")
            return "Error_Generating_PDF.pdf"

    def create_word(self, content, filename=None):
        """Generates a Microsoft Word (.docx) file."""
        if not filename:
            filename = f"doc_{uuid.uuid4().hex[:8]}.docx"
        
        if Document is None:
            return "Error: python-docx not installed"

        doc = Document()
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line: continue
            
            if line.startswith('# '):
                doc.add_heading(line.replace('# ', '').strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line.replace('## ', '').replace('### ', '').strip(), level=2)
            elif line.startswith('* ') or line.startswith('- '):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
            else:
                doc.add_paragraph(line)
                
        path = f"data/{filename}"
        doc.save(path)
        return path

    def create_excel(self, content, filename=None):
        """Generates an Excel (.xlsx) file."""
        if not filename:
            filename = f"sheet_{uuid.uuid4().hex[:8]}.xlsx"
        
        wb = openpyxl.Workbook()
        ws = wb.active
        
        rows = content.strip().split('\n')
        for row in rows:
            columns = [c.strip() for c in row.split(',')]
            ws.append(columns)
            
        path = f"data/{filename}"
        wb.save(path)
        return path


# from fpdf import FPDF
# import uuid
# import os
# import openpyxl
# from docx import Document  # Requires: pip install python-docx

# class DocumentTool:
#     def __init__(self):
#         if not os.path.exists("data"):
#             os.makedirs("data")

#     def create_pdf(self, content, filename=None):
#         """Generates a PDF with formatting."""
#         if not filename:
#             filename = f"doc_{uuid.uuid4().hex[:8]}.pdf"
            
#         pdf = FPDF()
#         pdf.add_page()
#         pdf.set_auto_page_break(auto=True, margin=15)
        
#         # Parse the content line by line
#         lines = content.split('\n')
        
#         for line in lines:
#             line = line.strip()
#             if not line:
#                 pdf.ln(3)  # Add small gap for empty lines
#                 continue

#             try:
#                 # CASE 1: Main Title (#)
#                 if line.startswith('# '):
#                     pdf.set_font("Arial", 'B', 16)
#                     clean_text = line.replace('# ', '').strip()
#                     pdf.multi_cell(0, 10, clean_text)
#                     pdf.ln(2)

#                 # CASE 2: Section Heading (## or ###)
#                 elif line.startswith('##') or line.startswith('###'):
#                     pdf.set_font("Arial", 'B', 12)
#                     clean_text = line.replace('### ', '').replace('## ', '').strip()
#                     pdf.multi_cell(0, 8, clean_text)
#                     pdf.ln(1)
                
#                 # CASE 3: Bullet Points (* or -)
#                 elif line.startswith('* ') or line.startswith('- '):
#                     pdf.set_font("Arial", '', 11)
#                     clean_text = line[2:].strip() # Remove the "* "
#                     current_x = pdf.get_x()
#                     pdf.set_x(current_x + 5) 
#                     pdf.multi_cell(0, 6, f"{chr(149)} {clean_text}")
#                     pdf.set_x(current_x) # Reset indent

#                 # CASE 4: Standard Paragraph Text
#                 else:
#                     pdf.set_font("Arial", '', 11)
#                     pdf.multi_cell(0, 6, line)
                    
#             except Exception as e:
#                 print(f"Error parsing line: {line} | {e}")
#                 pdf.set_font("Arial", '', 11)
#                 pdf.multi_cell(0, 6, line)

#         path = f"data/{filename}"
#         pdf.output(path)
#         return path

#     def create_word(self, content, filename=None):
#         """Generates a Microsoft Word (.docx) file."""
#         if not filename:
#             filename = f"doc_{uuid.uuid4().hex[:8]}.docx"
        
#         doc = Document()
#         lines = content.split('\n')
        
#         for line in lines:
#             line = line.strip()
#             if not line: continue
            
#             # Simple formatting mapping
#             if line.startswith('# '):
#                 doc.add_heading(line.replace('# ', '').strip(), level=1)
#             elif line.startswith('## '):
#                 doc.add_heading(line.replace('## ', '').replace('### ', '').strip(), level=2)
#             elif line.startswith('* ') or line.startswith('- '):
#                 doc.add_paragraph(line[2:].strip(), style='List Bullet')
#             else:
#                 doc.add_paragraph(line)
                
#         path = f"data/{filename}"
#         doc.save(path)
#         return path

#     def create_excel(self, content, filename=None):
#         """Generates an Excel (.xlsx) file from CSV-like text."""
#         if not filename:
#             filename = f"sheet_{uuid.uuid4().hex[:8]}.xlsx"
        
#         wb = openpyxl.Workbook()
#         ws = wb.active
        
#         # Assume LLM sends comma-separated lines
#         rows = content.strip().split('\n')
#         for row in rows:
#             # Clean and split by comma
#             columns = [c.strip() for c in row.split(',')]
#             ws.append(columns)
            
#         path = f"data/{filename}"
#         wb.save(path)
#         return path

# # from fpdf import FPDF
# # import uuid
# # import os

# # class DocumentTool:
# #     def __init__(self):
# #         if not os.path.exists("data"):
# #             os.makedirs("data")

# #     def create_pdf(self, content, filename=None):
# #         if not filename:
# #             filename = f"doc_{uuid.uuid4().hex[:8]}.pdf"
            
# #         pdf = FPDF()
# #         pdf.add_page()
# #         pdf.set_auto_page_break(auto=True, margin=15)
        
# #         # Parse the content line by line
# #         lines = content.split('\n')
        
# #         for line in lines:
# #             line = line.strip()
# #             if not line:
# #                 pdf.ln(3)  # Add small gap for empty lines
# #                 continue

# #             try:
# #                 # CASE 1: Main Title (#)
# #                 if line.startswith('# '):
# #                     pdf.set_font("Arial", 'B', 16)
# #                     # Clean the markdown symbol
# #                     clean_text = line.replace('# ', '').strip()
# #                     pdf.multi_cell(0, 10, clean_text)
# #                     pdf.ln(2)

# #                 # CASE 2: Section Heading (## or ###)
# #                 elif line.startswith('##') or line.startswith('###'):
# #                     pdf.set_font("Arial", 'B', 12)
# #                     clean_text = line.replace('### ', '').replace('## ', '').strip()
# #                     pdf.multi_cell(0, 8, clean_text)
# #                     pdf.ln(1)
                
# #                 # CASE 3: Bullet Points (* or -)
# #                 elif line.startswith('* ') or line.startswith('- '):
# #                     pdf.set_font("Arial", '', 11)
# #                     clean_text = line[2:].strip() # Remove the "* "
# #                     # Simulate a bullet point with indent
# #                     current_x = pdf.get_x()
# #                     pdf.set_x(current_x + 5) 
# #                     pdf.multi_cell(0, 6, f"{chr(149)} {clean_text}")
# #                     pdf.set_x(current_x) # Reset indent

# #                 # CASE 4: Standard Paragraph Text
# #                 else:
# #                     pdf.set_font("Arial", '', 11)
# #                     pdf.multi_cell(0, 6, line)
                    
# #             except Exception as e:
# #                 print(f"Error parsing line: {line} | {e}")
# #                 # Fallback ensures text isn't lost
# #                 pdf.set_font("Arial", '', 11)
# #                 pdf.multi_cell(0, 6, line)

# #         path = f"data/{filename}"
# #         pdf.output(path)
# #         return path